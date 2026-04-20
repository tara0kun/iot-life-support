"""家族会議用の企画書 (Word) を生成する。

実行: python scripts/build_plan_doc.py
出力: plan.docx (プロジェクトルート)
"""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent.parent / "plan.docx"

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
ACCENT = RGBColor(0xE8, 0x8B, 0x3B)
GRAY = RGBColor(0x40, 0x40, 0x40)

doc = Document()

# デフォルトフォント設定 (日本語)
style = doc.styles["Normal"]
style.font.name = "Meiryo"
style.font.size = Pt(10.5)
rpr = style.element.get_or_add_rPr()
rfonts = rpr.find(qn("w:rFonts"))
if rfonts is None:
    from docx.oxml import OxmlElement
    rfonts = OxmlElement("w:rFonts")
    rpr.append(rfonts)
rfonts.set(qn("w:eastAsia"), "Meiryo")

# ページ余白を少し広めに
for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)


def h(text, level=1, color=NAVY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.bold = True
    run.font.color.rgb = color
    sizes = {1: 20, 2: 14, 3: 12}
    run.font.size = Pt(sizes.get(level, 11))
    run.font.name = "Meiryo"
    rpr = run._r.get_or_add_rPr()
    from docx.oxml import OxmlElement
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), "Meiryo")
    rpr.append(rfonts)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
    else:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    return p


def para(text, *, bold=False, size=10.5, color=None, align=None):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.name = "Meiryo"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    from docx.oxml import OxmlElement
    rpr = run._r.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), "Meiryo")
    rpr.append(rfonts)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + 0.5 * level)
    run = p.add_run(text)
    run.font.name = "Meiryo"
    run.font.size = Pt(10.5)
    from docx.oxml import OxmlElement
    rpr = run._r.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), "Meiryo")
    rpr.append(rfonts)
    return p


def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, hv in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(hv)
        r.font.bold = True
        r.font.size = Pt(10.5)
        r.font.name = "Meiryo"
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        from docx.oxml import OxmlElement
        rpr = r._r.get_or_add_rPr()
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:eastAsia"), "Meiryo")
        rpr.append(rfonts)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(v))
            r.font.size = Pt(10)
            r.font.name = "Meiryo"
            from docx.oxml import OxmlElement
            rpr = r._r.get_or_add_rPr()
            rfonts = OxmlElement("w:rFonts")
            rfonts.set(qn("w:eastAsia"), "Meiryo")
            rpr.append(rfonts)
    if col_widths:
        for col, w in zip(t.columns, col_widths):
            for c in col.cells:
                c.width = Cm(w)


# ============================================================
# 表紙
# ============================================================
para("企画書", bold=True, size=14, color=GRAY, align="right")
para(f"作成日: {date.today().isoformat()}", size=10, color=GRAY, align="right")
doc.add_paragraph()
doc.add_paragraph()

para("認知症の祖母のための", bold=True, size=22, color=NAVY, align="center")
para("IoT生活サポートシステム", bold=True, size=26, color=NAVY, align="center")
para("導入企画書", bold=True, size=18, color=ACCENT, align="center")

for _ in range(3):
    doc.add_paragraph()

para("── 祖母の尊厳と安心、家族の負担軽減を両立するために ──",
     size=12, color=GRAY, align="center")

for _ in range(6):
    doc.add_paragraph()

para("提案者: 孫", size=11, align="right")
para("対象: 家族（母・祖父・本人を含む家族全員）", size=11, align="right")
doc.add_page_break()

# ============================================================
# 1. 企画概要
# ============================================================
h("1. 企画概要")

h("1.1 背景", 2)
para(
    "祖母の認知症が進行し、日常生活において以下の課題が顕在化しています。"
    "家族は可能な限り支援したいと考えていますが、24時間の付き添いは現実的ではなく、"
    "また直接的な指摘は祖母の尊厳を傷つけ、関係性の悪化を招くジレンマを抱えています。"
)

h("1.2 目的", 2)
para(
    "IoT（モノのインターネット）技術を活用し、祖母が自分らしく生活できる時間を"
    "できる限り長く保ちながら、家族の精神的・身体的負担を軽減することを目的とします。"
)

h("1.3 基本方針", 2)
bullet("祖母を「監視」するのではなく、本人が自分の行動を確認できる「記録帳」として提示する")
bullet("機械の不調に見せることで、祖母を傷つけずに危険行動を物理的に阻止する")
bullet("家族による操作・編集は、祖母に絶対に知られないよう完全に分離する")
bullet("技術で置き換えるのではなく、家族のケアを補助するツールとして位置づける")

# ============================================================
# 2. 解決すべき課題
# ============================================================
h("2. 解決すべき課題")

h("2.1 食事に関する課題（最優先）", 2)
bullet("朝食を摂ったことを忘れ、数分後にまた炊飯器を開けてしまう")
bullet("冷蔵庫から何度もおかずを取り出して食べ、健康を損なう恐れがある")
bullet("生米の上にご飯を入れて再炊飯するなど、器具の誤使用が発生")
bullet("指摘すると強く反発し、関係が険悪になる")

h("2.2 訪問者への対応", 2)
bullet("訪問販売や勧誘に応じて、不要な契約を結んでしまう")
bullet("同じ販売員が何度来ているか把握できない")
bullet("来訪者の内容を家族が後から確認する手段がない")

h("2.3 入浴時の衛生管理", 2)
bullet("お風呂に入っても頭を洗わない日がある")
bullet("「洗った」と本人は主張するため、家族が促すと険悪になる")

# ============================================================
# 3. 提案するシステム
# ============================================================
h("3. 提案するシステムの概要")

h("3.1 全体アーキテクチャ", 2)
para(
    "Raspberry Pi（小型コンピュータ）を中央サーバとし、スマートプラグ・"
    "開閉センサー・小型カメラを自宅に設置。祖母の行動を自動で記録し、"
    "以下3つのタイミングで介入します。"
)
table(
    ["段階", "介入の内容", "祖母の体感"],
    [
        ["第1段階", "リビングのタブレットに「今日の記録」を表示",
         "「あら、もう食べたんだ」と自分で気づける"],
        ["第2段階", "炊飯器・IHの電源が一時的に入らなくなる",
         "「調子が悪いのかしら」と思うだけ"],
        ["第3段階", "家族のLINEに通知が届き、さりげなく声かけ",
         "家族との普段どおりの会話の一部"],
    ],
    col_widths=[2.5, 6, 7],
)

h("3.2 宣言型アンロックと顔認識の組み合わせ", 2)
para(
    "冷蔵庫・炊飯器・IH等の食料源は常時ロック状態とし、以下のいずれかの方法で"
    "解錠します。これにより祖母・家族の双方に自然な操作体験を提供します。"
)
bullet("カメラによる顔認識：近づいた人物を自動で識別し、必要に応じて警告を出しつつ解錠")
bullet("タブレットでの宣言：カメラで識別できない時はボタン操作で「誰が使うか」を選択")
bullet("家族がアクセスした場合は、無言で即解錠され、祖母の画面には一切記録されない")

h("3.3 多層防御モデル", 2)
para("単一の対策に頼らず、以下の4層で段階的に安全を確保します。")
table(
    ["層", "役割", "手段"],
    [
        ["Layer 1", "気づかせる", "タブレットの記録帳表示"],
        ["Layer 2", "物理的に阻止する", "スマートプラグ・スマートロックによる宣言型アンロック"],
        ["Layer 3", "家族に通知する", "LINE Messaging APIで自然な声かけを誘導"],
        ["Layer 4", "環境を整える", "祖母用プレートの運用・物理分離等（家族の運用）"],
    ],
    col_widths=[2.5, 4, 9],
)

# ============================================================
# 4. プライバシー・倫理配慮
# ============================================================
h("4. プライバシー・倫理配慮")

h("4.1 祖母への配慮", 2)
bullet("カメラの映像そのものは保存せず、『いつ・誰が・何をした』という記録のみをローカルDBに残す")
bullet("祖母の画面には、家族の行動や編集操作は一切表示されない")
bullet("祖母が不審に思うきっかけを最小化する（機械は静か、録画ランプなし等）")
bullet("システムの存在をどこまで伝えるかは家族で合意し、本人の状態に応じて調整する")

h("4.2 家族への配慮", 2)
bullet("家族メンバー各自の行動も個別に記録されるが、本人のみ閲覧・修正可能")
bullet("祖父への同意取得が必須（映像に映り込む可能性があるため）")
bullet("家族間で誤った記録は相互に修正でき、祖母には変更痕跡が見えない")

h("4.3 データセキュリティ", 2)
bullet("すべてのデータは自宅のRaspberry Pi内に保存される（クラウド送信なし）")
bullet("家族が外部から閲覧する場合は、VPN等で暗号化された経路を使う")
bullet("データの保存期間は家族が設定可能（例: 90日で自動削除）")

# ============================================================
# 5. 必要な機材と予算
# ============================================================
h("5. 必要な機材と予算")

h("5.1 機材一覧（段階別）", 2)
table(
    ["段階", "機材", "数量", "概算金額"],
    [
        ["お試し", "Tapo P115（スマートプラグ）", "1", "約 2,000円"],
        ["お試し", "Tapo T110（開閉センサー）", "2", "約 5,000円"],
        ["お試し", "Tapo H100（ハブ）", "1", "約 3,500円"],
        ["", "お試し段階 小計", "", "約 10,500円"],
        ["本格導入", "Tapo T110（追加）", "2", "約 5,000円"],
        ["本格導入", "Tapo C210（カメラ）", "1", "約 4,000円"],
        ["本格導入", "Tapo P115（追加）", "1", "約 2,000円"],
        ["本格導入", "タブレット端末", "1", "既存流用 or 15,000〜30,000円"],
        ["", "本格導入段階 小計（新規購入時）", "", "約 30,000円〜45,000円"],
        ["拡張", "Tapo D230S1（玄関ドアホン）", "1", "約 15,000円"],
        ["拡張", "SwitchBot Lock（冷蔵庫用）", "1〜2", "約 9,000〜18,000円"],
        ["拡張", "スマートスピーカー（任意）", "1", "約 5,000〜15,000円"],
        ["", "拡張段階 小計", "", "約 30,000〜60,000円"],
        ["", "【合計（全機能導入時）】", "", "約 70,000〜120,000円"],
    ],
    col_widths=[3, 7, 2, 4],
)
para("※Raspberry Pi本体（中央サーバ）は既に所有しているため追加費用不要。",
     size=9.5, color=GRAY)

h("5.2 その他コスト", 2)
bullet("LINE Messaging API：月1,000通まで無料。通常利用なら追加費用なし")
bullet("電気代の増加：スマートプラグ類は微々たるもの（月数十円程度）")
bullet("通信費：既存の自宅Wi-Fiを使用するため追加なし")

# ============================================================
# 6. 導入スケジュール
# ============================================================
h("6. 導入スケジュール")

table(
    ["時期", "段階", "内容"],
    [
        ["今", "家族会議", "本企画書を元に、家族全員で合意形成"],
        ["1〜2週間", "自宅試験", "提案者（孫）の自宅で動作検証"],
        ["1ヶ月目", "お試し導入", "炊飯器のみを対象に実家に設置、祖母の反応を観察"],
        ["2〜3ヶ月目", "本格導入", "効果があれば冷蔵庫・IH・カメラへ拡張"],
        ["3ヶ月目以降", "調整・運用", "数値を見ながら介入強度を調整、他の問題にも対応"],
    ],
    col_widths=[3, 3, 10],
)

# ============================================================
# 7. 体制と役割分担
# ============================================================
h("7. 体制と役割分担（案）")

table(
    ["役割", "担当者（案）", "内容"],
    [
        ["システム開発・保守", "孫", "コード開発、機器設置、トラブル対応"],
        ["日常の確認・修正", "母", "家族UIで記録確認、誤記録の修正"],
        ["LINE通知の対応", "母・祖父", "祖母への声かけ"],
        ["物理機器のメンテ", "祖父", "機器の電池交換、設置位置調整"],
        ["最終決定", "家族全員", "導入可否、拡張方針、中断判断"],
    ],
    col_widths=[4, 3, 9],
)

# ============================================================
# 8. リスクと対策
# ============================================================
h("8. リスクと対策")

table(
    ["リスク", "対策"],
    [
        ["祖母が異変に気づき拒否反応を示す",
         "設計原則を厳守。うまくいかなければ速やかに撤去する柔軟性を持つ"],
        ["カメラの誤認識により家族の行動が祖母の記録になる",
         "家族UIからワンタップで修正可能。修正操作は祖母に見えない"],
        ["機器の故障によりロックが解除できなくなる",
         "家族スマホからの緊急解除機能あり。故障時は「常時開」側に倒れる設計"],
        ["インターネットが切断される",
         "記録機能はローカルで継続。LINE通知は復旧後に配信"],
        ["システム由来のトラブルが発生する",
         "月次で家族レビュー。中断基準を事前に定義"],
        ["祖母の状態悪化により、このシステムでは対応できなくなる",
         "システムはあくまで補助。医療・介護サービスを主とし、本システムは補完にとどめる"],
    ],
    col_widths=[7, 9],
)

# ============================================================
# 9. 期待される効果
# ============================================================
h("9. 期待される効果")

h("9.1 祖母にとって", 2)
bullet("食べ過ぎによる健康悪化のリスクを低減")
bullet("家族から直接指摘される機会が減り、尊厳と穏やかさを保てる")
bullet("訪問販売等の被害を未然に防ぐ")
bullet("できる限り長く自宅で生活を続けられる可能性が高まる")

h("9.2 家族にとって", 2)
bullet("24時間の注意監視から解放され、精神的負担が軽減")
bullet("「指摘するか、見逃すか」のジレンマから解放される")
bullet("外出中でも状況を確認でき、緊急時のみ介入すれば良い")
bullet("ケアマネージャー・主治医への状況共有が容易になる（記録データがあるため）")

# ============================================================
# 10. 家族に決めていただきたいこと
# ============================================================
h("10. 家族に決めていただきたいこと")
para("次回の家族会議までに、以下について話し合ってください。")

table(
    ["項目", "選択肢"],
    [
        ["取り組む問題の優先順位",
         "食事／訪問者／入浴 のうちどれから始めるか"],
        ["初期投資の予算",
         "お試し（1万円）／本格導入（5万円）／フル装備（12万円）"],
        ["祖母への説明方針",
         "完全に伏せる／「記録帳」として緩やかに伝える／全て説明する"],
        ["祖父・家族の同意",
         "カメラ設置・映像処理への同意"],
        ["日常の役割分担",
         "通知対応・記録修正を誰が担当するか"],
        ["中断・撤去の基準",
         "どのような事態になれば中断するか（例: 祖母が嫌がった時）"],
        ["導入タイミング",
         "いつから自宅試験・実家投入を始めるか"],
    ],
    col_widths=[5, 11],
)

# ============================================================
# 11. 結び
# ============================================================
h("11. 結び")
para(
    "このシステムは、祖母の日々を「技術で置き換える」ことを目的としていません。"
    "家族のやさしさが、少しでも続けやすくなるための、静かな補助装置です。"
)
para(
    "おばあちゃんができるだけ長く、自分らしく暮らせるように。"
    "そして家族が疲弊しすぎずに支え続けられるように。"
)
para(
    "本企画書の内容について、ご意見・ご質問・ご懸念があれば、"
    "どのような些細なものでも遠慮なくお聞かせください。"
    "家族全員が納得できる形で進めることを何より大切にしたいと考えています。"
)

doc.add_paragraph()
doc.add_paragraph()
para("── 以上 ──", color=GRAY, align="center")

doc.save(OUT)
print(f"作成完了: {OUT}")
print(f"サイズ: {OUT.stat().st_size // 1024} KB")
